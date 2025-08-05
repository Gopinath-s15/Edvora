"""
Document Processing Utilities for Edvora
Handles document downloading, parsing, and text extraction
"""

import os
import tempfile
import logging
from typing import Optional, Dict, Any
from pathlib import Path
import httpx
import aiofiles
from PyPDF2 import PdfReader
from docx import Document
import re

logger = logging.getLogger(__name__)

class DocumentProcessor:
    """Handles document downloading and text extraction"""
    
    def __init__(self):
        self.max_file_size = int(os.getenv("MAX_FILE_SIZE_MB", 50)) * 1024 * 1024  # Convert to bytes
        self.supported_extensions = {'.pdf', '.docx'}
        
    async def process_document_from_url(self, url: str) -> str:
        """
        Download and process document from URL
        
        Args:
            url: URL to the document
            
        Returns:
            Extracted text content
        """
        try:
            logger.info(f"Starting document processing for URL: {url}")
            
            # Download document
            file_path = await self._download_document(url)
            
            try:
                # Extract text based on file type
                text = await self._extract_text(file_path)
                
                # Clean and validate text
                cleaned_text = self._clean_text(text)
                
                if not cleaned_text:
                    raise ValueError("No text content extracted from document")
                
                logger.info(f"Successfully extracted {len(cleaned_text)} characters from document")
                return cleaned_text
                
            finally:
                # Clean up temporary file
                if file_path and os.path.exists(file_path):
                    os.unlink(file_path)
                    
        except Exception as e:
            logger.error(f"Error processing document from URL {url}: {str(e)}")
            raise
    
    async def _download_document(self, url: str) -> str:
        """Download document from URL to temporary file"""
        try:
            logger.info(f"Downloading document from: {url}")
            
            async with httpx.AsyncClient(timeout=30.0, follow_redirects=True) as client:
                response = await client.get(url)
                response.raise_for_status()
                
                # Check content type and size
                content_type = response.headers.get('content-type', '').lower()
                content_length = int(response.headers.get('content-length', 0))
                
                if content_length > self.max_file_size:
                    raise ValueError(f"File size ({content_length} bytes) exceeds maximum allowed size")
                
                # Determine file extension from URL or content type
                file_extension = self._get_file_extension(url, content_type)
                
                # Create temporary file
                with tempfile.NamedTemporaryFile(delete=False, suffix=file_extension) as temp_file:
                    temp_file.write(response.content)
                    temp_file_path = temp_file.name
                
                logger.info(f"Document downloaded to: {temp_file_path}")
                return temp_file_path
                
        except httpx.HTTPError as e:
            logger.error(f"HTTP error downloading document: {str(e)}")
            raise ValueError(f"Failed to download document: {str(e)}")
        except Exception as e:
            logger.error(f"Error downloading document: {str(e)}")
            raise
    
    def _get_file_extension(self, url: str, content_type: str) -> str:
        """Determine file extension from URL or content type"""
        # Try to get extension from URL
        url_path = Path(url.split('?')[0])  # Remove query parameters
        if url_path.suffix.lower() in self.supported_extensions:
            return url_path.suffix.lower()
        
        # Try to determine from content type
        content_type_mapping = {
            'application/pdf': '.pdf',
            'application/vnd.openxmlformats-officedocument.wordprocessingml.document': '.docx',
            'application/msword': '.docx'
        }
        
        for ct, ext in content_type_mapping.items():
            if ct in content_type:
                return ext
        
        # Default to PDF if uncertain
        logger.warning(f"Could not determine file type from URL or content-type, defaulting to PDF")
        return '.pdf'
    
    async def _extract_text(self, file_path: str) -> str:
        """Extract text from document based on file type"""
        file_extension = Path(file_path).suffix.lower()
        
        if file_extension == '.pdf':
            return await self._extract_pdf_text(file_path)
        elif file_extension == '.docx':
            return await self._extract_docx_text(file_path)
        else:
            raise ValueError(f"Unsupported file type: {file_extension}")
    
    async def _extract_pdf_text(self, file_path: str) -> str:
        """Extract text from PDF file"""
        try:
            text = ""
            with open(file_path, 'rb') as file:
                pdf_reader = PdfReader(file)
                
                for page_num, page in enumerate(pdf_reader.pages):
                    try:
                        page_text = page.extract_text()
                        if page_text:
                            text += f"\n--- Page {page_num + 1} ---\n{page_text}\n"
                    except Exception as e:
                        logger.warning(f"Error extracting text from page {page_num + 1}: {str(e)}")
                        continue
            
            if not text.strip():
                raise ValueError("No text could be extracted from PDF")
            
            return text
            
        except Exception as e:
            logger.error(f"Error extracting PDF text: {str(e)}")
            raise ValueError(f"Failed to extract text from PDF: {str(e)}")
    
    async def _extract_docx_text(self, file_path: str) -> str:
        """Extract text from DOCX file"""
        try:
            doc = Document(file_path)
            text = ""
            
            # Extract text from paragraphs
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text += paragraph.text + "\n"
            
            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            text += cell.text + " "
                    text += "\n"
            
            if not text.strip():
                raise ValueError("No text could be extracted from DOCX")
            
            return text
            
        except Exception as e:
            logger.error(f"Error extracting DOCX text: {str(e)}")
            raise ValueError(f"Failed to extract text from DOCX: {str(e)}")
    
    def _clean_text(self, text: str) -> str:
        """Clean and normalize extracted text"""
        if not text:
            return ""
        
        # Remove excessive whitespace
        text = re.sub(r'\s+', ' ', text)
        
        # Remove special characters that might interfere with processing
        text = re.sub(r'[^\w\s\.\,\;\:\!\?\-\(\)\[\]\{\}\"\'₹\$\%\@\#\&\*\+\=\<\>\/\\]', ' ', text)
        
        # Normalize line breaks
        text = re.sub(r'\n+', '\n', text)
        
        # Remove leading/trailing whitespace
        text = text.strip()
        
        return text
    
    def get_document_metadata(self, file_path: str) -> Dict[str, Any]:
        """Extract metadata from document"""
        try:
            file_stats = os.stat(file_path)
            file_extension = Path(file_path).suffix.lower()
            
            metadata = {
                'file_size': file_stats.st_size,
                'file_type': file_extension,
                'created_time': file_stats.st_ctime,
                'modified_time': file_stats.st_mtime
            }
            
            # Add document-specific metadata
            if file_extension == '.pdf':
                try:
                    with open(file_path, 'rb') as file:
                        pdf_reader = PdfReader(file)
                        metadata.update({
                            'page_count': len(pdf_reader.pages),
                            'pdf_info': pdf_reader.metadata if pdf_reader.metadata else {}
                        })
                except:
                    pass
            
            elif file_extension == '.docx':
                try:
                    doc = Document(file_path)
                    metadata.update({
                        'paragraph_count': len(doc.paragraphs),
                        'table_count': len(doc.tables)
                    })
                except:
                    pass
            
            return metadata
            
        except Exception as e:
            logger.warning(f"Could not extract metadata: {str(e)}")
            return {}
